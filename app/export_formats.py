"""
export_formats.py  –  Export resume to DOCX, plain text, and Markdown
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _get_profile_data(profile_data: dict) -> dict:
    """Normalize profile data."""
    skills = profile_data.get("skills", {})
    if hasattr(skills, "__dict__"):
        skills = vars(skills)
    return {
        "basics": profile_data.get("basics", {}),
        "education": profile_data.get("education", []),
        "experience": profile_data.get("experience", []),
        "projects": profile_data.get("projects", []),
        "skills": skills,
        "certifications": profile_data.get("certifications", []),
        "awards": profile_data.get("awards", []),
        "publications": profile_data.get("publications", []),
    }


def export_to_docx(profile_data: dict, output_path: str) -> str:
    """Export resume as a .docx file."""
    d = _get_profile_data(profile_data)
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)

    b = d["basics"]

    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name_para.add_run(b.get("name", ""))
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(27, 58, 92)

    contact_parts = []
    if b.get("phone"): contact_parts.append(b["phone"])
    if b.get("email"): contact_parts.append(b["email"])
    if b.get("linkedin"): contact_parts.append(b["linkedin"])
    if b.get("github"): contact_parts.append(b["github"])

    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = contact_para.add_run(" | ".join(contact_parts))
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(102, 102, 102)

    def add_section_header(text: str):
        h = doc.add_paragraph()
        run = h.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(11)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(2)
        doc.add_paragraph().add_run("_" * 70).font.color.rgb = RGBColor(200, 200, 200)

    def add_two_col(left: str, right: str):
        p = doc.add_paragraph()
        run_l = p.add_run(left)
        run_l.bold = True
        run_l.font.size = Pt(10)
        tab_stops = p.paragraph_format.tab_stops
        p.add_run("\t" * 6 + right).font.size = Pt(9)
        p.paragraph_format.space_after = Pt(1)

    def add_bullets(items: list):
        for item in items:
            p = doc.add_paragraph(item, style="List Bullet")
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(1)
            for run in p.runs:
                run.font.size = Pt(9)

    if d["education"]:
        add_section_header("Education")
        for edu in d["education"]:
            add_two_col(edu.get("institution", ""), f'{edu.get("startDate","")} - {edu.get("endDate","")}')
            doc.add_paragraph(edu.get("degree", "")).runs[0].italic = True
            doc.paragraphs[-1].runs[0].font.size = Pt(9)

    if d["experience"]:
        add_section_header("Professional Experience")
        for exp in d["experience"]:
            add_two_col(exp.get("role", ""), f'{exp.get("startDate","")} - {exp.get("endDate","")}')
            doc.add_paragraph(exp.get("company", "")).runs[0].italic = True
            doc.paragraphs[-1].runs[0].font.size = Pt(9)
            if exp.get("bullets"):
                add_bullets(exp["bullets"])

    if d["projects"]:
        add_section_header("Projects")
        for proj in d["projects"]:
            techs = " - ".join(proj.get("technologies", []))
            p = doc.add_paragraph()
            run = p.add_run(proj.get("name", ""))
            run.bold = True
            run.font.size = Pt(10)
            p.add_run(f"  |  {techs}").font.size = Pt(8)
            p.add_run(f"  |  {techs}").font.size = Pt(8)
            p.add_run(f"  |  {techs}").font.size = Pt(8)
            p.add_run(f"  |  {techs}").font.size = Pt(8)
            if proj.get("bullets"):
                add_bullets(proj["bullets"])

    if d["skills"]:
        add_section_header("Technical Skills")
        skill_parts = []
        for label, key in [("Languages", "languages"), ("Frameworks", "frameworks"), ("Tools & Platforms", "tools")]:
            items = d["skills"].get(key, [])
            if items:
                skill_parts.append(f"{label}: {', '.join(items)}")
        doc.add_paragraph(" | ".join(skill_parts))

    if d["certifications"]:
        add_section_header("Certifications")
        for cert in d["certifications"]:
            text = f'{cert.get("name","")}'
            if cert.get("issuer"): text += f' - {cert["issuer"]}'
            if cert.get("date"): text += f' ({cert["date"]})'
            doc.add_paragraph(text)

    if d["awards"]:
        add_section_header("Awards & Honors")
        for award in d["awards"]:
            text = f'{award.get("title","")}'
            if award.get("organization"): text += f' - {award["organization"]}'
            if award.get("date"): text += f' ({award["date"]})'
            doc.add_paragraph(text)

    if d["publications"]:
        add_section_header("Publications")
        for pub in d["publications"]:
            text = f'{pub.get("title","")}'
            if pub.get("venue"): text += f'. {pub["venue"]}'
            if pub.get("date"): text += f', {pub["date"]}'
            doc.add_paragraph(text)

    os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else ".", exist_ok=True)
    doc.save(output_path)
    return output_path


def export_to_text(profile_data: dict, output_path: str) -> str:
    """Export resume as plain text."""
    d = _get_profile_data(profile_data)
    lines = []
    b = d["basics"]

    lines.append("=" * 70)
    lines.append(b.get("name", "").upper())
    lines.append("=" * 70)

    contact_parts = []
    if b.get("phone"): contact_parts.append(b["phone"])
    if b.get("email"): contact_parts.append(b["email"])
    if b.get("linkedin"): contact_parts.append(b["linkedin"])
    if b.get("github"): contact_parts.append(b["github"])
    if contact_parts:
        lines.append(" | ".join(contact_parts))
    lines.append("")

    if d["education"]:
        lines.append("-" * 70)
        lines.append("EDUCATION")
        lines.append("-" * 70)
        for edu in d["education"]:
            lines.append(f'{edu.get("institution","")}  |  {edu.get("startDate","")} - {edu.get("endDate","")}')
            lines.append(f'  {edu.get("degree","")}  |  {edu.get("location","")}')
            lines.append("")

    if d["experience"]:
        lines.append("-" * 70)
        lines.append("PROFESSIONAL EXPERIENCE")
        lines.append("-" * 70)
        for exp in d["experience"]:
            lines.append(f'{exp.get("role","")}  |  {exp.get("startDate","")} - {exp.get("endDate","")}')
            lines.append(f'  {exp.get("company","")}  |  {exp.get("location","")}')
            for bullet in exp.get("bullets", []):
                lines.append(f"  * {bullet}")
            lines.append("")

    if d["projects"]:
        lines.append("-" * 70)
        lines.append("PROJECTS")
        lines.append("-" * 70)
        for proj in d["projects"]:
            techs = ", ".join(proj.get("technologies", []))
            lines.append(f'{proj.get("name","")}  |  {techs}')
            for bullet in proj.get("bullets", []):
                lines.append(f"  * {bullet}")
            lines.append("")

    if d["skills"]:
        lines.append("-" * 70)
        lines.append("TECHNICAL SKILLS")
        lines.append("-" * 70)
        for label, key in [("Languages", "languages"), ("Frameworks", "frameworks"), ("Tools & Platforms", "tools")]:
            items = d["skills"].get(key, [])
            if items:
                lines.append(f"{label}: {', '.join(items)}")
        lines.append("")

    if d["certifications"]:
        lines.append("-" * 70)
        lines.append("CERTIFICATIONS")
        lines.append("-" * 70)
        for cert in d["certifications"]:
            text = cert.get("name", "")
            if cert.get("issuer"): text += f' - {cert["issuer"]}'
            if cert.get("date"): text += f' ({cert["date"]})'
            lines.append(text)
        lines.append("")

    if d["awards"]:
        lines.append("-" * 70)
        lines.append("AWARDS & HONORS")
        lines.append("-" * 70)
        for award in d["awards"]:
            text = award.get("title", "")
            if award.get("organization"): text += f' - {award["organization"]}'
            if award.get("date"): text += f' ({award["date"]})'
            lines.append(text)
        lines.append("")

    if d["publications"]:
        lines.append("-" * 70)
        lines.append("PUBLICATIONS")
        lines.append("-" * 70)
        for pub in d["publications"]:
            text = pub.get("title", "")
            if pub.get("venue"): text += f'. {pub["venue"]}'
            if pub.get("date"): text += f', {pub["date"]}'
            lines.append(text)
        lines.append("")

    os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else ".", exist_ok=True)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    return output_path


def export_to_markdown(profile_data: dict, output_path: str) -> str:
    """Export resume as Markdown."""
    d = _get_profile_data(profile_data)
    lines = []
    b = d["basics"]

    lines.append(f'# {b.get("name", "")}')
    lines.append("")

    contact_parts = []
    if b.get("phone"): contact_parts.append(b["phone"])
    if b.get("email"): contact_parts.append(f'[{b["email"]}](mailto:{b["email"]})')
    if b.get("linkedin"): contact_parts.append(f'[{b["linkedin"]}](https://{b["linkedin"]})')
    if b.get("github"): contact_parts.append(f'[{b["github"]}](https://{b["github"]})')
    if contact_parts:
        lines.append(" | ".join(contact_parts))
    lines.append("")

    if d["education"]:
        lines.append("## Education")
        lines.append("")
        for edu in d["education"]:
            lines.append(f'**{edu.get("institution","")}**  ')
            lines.append(f'{edu.get("startDate","")} - {edu.get("endDate","")}  ')
            lines.append(f'*{edu.get("degree","")}*  ')
            lines.append(f'{edu.get("location","")}')
            lines.append("")

    if d["experience"]:
        lines.append("## Professional Experience")
        lines.append("")
        for exp in d["experience"]:
            lines.append(f'**{exp.get("role","")}**  ')
            lines.append(f'{exp.get("startDate","")} - {exp.get("endDate","")}  ')
            lines.append(f'*{exp.get("company","")}* | {exp.get("location","")}')
            lines.append("")
            for bullet in exp.get("bullets", []):
                lines.append(f"- {bullet}")
            lines.append("")

    if d["projects"]:
        lines.append("## Projects")
        lines.append("")
        for proj in d["projects"]:
            techs = ", ".join(proj.get("technologies", []))
            lines.append(f'**{proj.get("name","")}** | *{techs}*')
            lines.append("")
            for bullet in proj.get("bullets", []):
                lines.append(f"- {bullet}")
            lines.append("")

    if d["skills"]:
        lines.append("## Technical Skills")
        lines.append("")
        for label, key in [("Languages", "languages"), ("Frameworks", "frameworks"), ("Tools & Platforms", "tools")]:
            items = d["skills"].get(key, [])
            if items:
                lines.append(f"- **{label}:** {', '.join(items)}")
        lines.append("")

    if d["certifications"]:
        lines.append("## Certifications")
        lines.append("")
        for cert in d["certifications"]:
            text = f'- **{cert.get("name","")}'
            if cert.get("issuer"): text += f' - {cert["issuer"]}'
            if cert.get("date"): text += f' ({cert["date"]})'
            text += "**"
            lines.append(text)
        lines.append("")

    if d["awards"]:
        lines.append("## Awards & Honors")
        lines.append("")
        for award in d["awards"]:
            text = f'- **{award.get("title","")}'
            if award.get("organization"): text += f' - {award["organization"]}'
            if award.get("date"): text += f' ({award["date"]})'
            text += "**"
            lines.append(text)
        lines.append("")

    if d["publications"]:
        lines.append("## Publications")
        lines.append("")
        for pub in d["publications"]:
            text = f'- **{pub.get("title","")}'
            if pub.get("venue"): text += f'. *{pub["venue"]}*'
            if pub.get("date"): text += f', {pub["date"]}'
            text += "**"
            lines.append(text)
        lines.append("")

    os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else ".", exist_ok=True)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    return output_path
